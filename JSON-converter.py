import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert  # Windows-only
from ebooklib import epub
import os

# === CSS for the beautiful HTML output ===
HTML_CSS_STYLE = """
<style>
    body {
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
        line-height: 1.6;
        color: #333;
        background-color: #f8f9fa;
        margin: 0;
        padding: 2rem;
    }
    .container {
        max-width: 800px;
        margin: 0 auto;
        background-color: #ffffff;
        border-radius: 10px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
        overflow: hidden;
    }
    .entry {
        padding: 2rem 2.5rem;
        border-bottom: 1px solid #e9ecef;
    }
    .entry:last-child {
        border-bottom: none;
    }
    .entry-date {
        font-size: 1.1rem;
        font-weight: 600;
        color: #007bff;
        margin-bottom: 1rem;
    }
    .entry-content h1, .entry-content h2, .entry-content h3, .entry-content h4, .entry-content h5, .entry-content h6 {
        color: #495057;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
        line-height: 1.3;
    }
    .entry-content h1 { font-size: 1.8em; }
    .entry-content h2 { font-size: 1.5em; }
    .entry-content h3 { font-size: 1.25em; }
    .entry-content p {
        margin-top: 0;
        margin-bottom: 1rem;
    }
    .entry-content a {
        color: #0056b3;
        text-decoration: none;
    }
    .entry-content a:hover {
        text-decoration: underline;
    }
    .entry-content code {
        background-color: #e9ecef;
        padding: 0.2em 0.4em;
        margin: 0;
        font-size: 85%;
        border-radius: 3px;
    }
    .entry-content pre {
        background-color: #e9ecef;
        padding: 1rem;
        border-radius: 5px;
        overflow-x: auto;
    }
    .entry-content pre code {
        padding: 0;
        font-size: inherit;
        color: inherit;
        background-color: transparent;
    }
    .main-title {
        text-align: center;
        padding: 2rem;
        background-color: #007bff;
        color: white;
    }
    .main-title h1 {
        margin: 0;
        font-size: 2.5rem;
    }
</style>
"""

# === Configure Input File ===
input_filename = "Journal1.json" 

# === Setup output folder and prefix ===
folder_name = os.path.splitext(input_filename)[0]
os.makedirs(folder_name, exist_ok=True)
today_str = datetime.now().strftime("%Y-%m-%d")
base_filename = f"output_{os.path.basename(input_filename)}_{today_str}"
output_prefix = os.path.join(folder_name, base_filename)

# === Process notes based on file type ===
notes = []
if input_filename.lower().endswith('.json'):
    print(f"Processing Day One JSON file: {input_filename}")
    with open(input_filename, "r", encoding="utf-8") as f: data = json.load(f)
    entries = data.get("entries", [])
    if not entries:
        print("[!] Error: No 'entries' found in the JSON file."); exit()
    entries.sort(key=lambda x: x.get('creationDate', ''))
    for entry in entries:
        date_obj = datetime.strptime(entry.get("creationDate", ""), "%Y-%m-%dT%H:%M:%SZ")
        notes.append({'date': date_obj.strftime("%Y-%m-%d"), 'text': entry.get("text", "").strip()})
elif input_filename.lower().endswith('.md'):
    print(f"Processing Markdown file: {input_filename}")
    with open(input_filename, "r", encoding="utf-8") as f: md_content = f.read()
    mod_time = os.path.getmtime(input_filename)
    mod_date_obj = datetime.fromtimestamp(mod_time)
    notes.append({'date': mod_date_obj.strftime("%Y-%m-%d"), 'text': md_content.strip()})
else:
    print(f"[!] Error: Unsupported file type for '{input_filename}'. Please use a .json or .md file."); exit()

# === Markdown Processing Functions ===
def markdown_to_plain_text(text):
    text = re.sub(r'^#{1,6}\s+(.+)$', r'\1', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*|__(.+?)__', r'\1\2', text)
    text = re.sub(r'\*(.+?)\*|_(.+?)_', r'\1\2', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'```.*?\n(.*?)\n```', r'\1', text, flags=re.DOTALL)
    return text

def markdown_to_latex(text):
    text = re.sub(r'^# (.+)$', r'\\section{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'\\subsection{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.+)$', r'\\subsubsection{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*|__(.+?)__', r'\\textbf{\1\2}', text)
    text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)|(?<!_)_([^_]+?)_(?!_)', r'\\textit{\1\2}', text)
    text = re.sub(r'`(.+?)`', r'\\texttt{\1}', text)
    text = re.sub(r'```.*?\n(.*?)\n```', r'\\begin{verbatim}\n\1\n\\end{verbatim}', text, flags=re.DOTALL)
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'\\href{\2}{\1}', text)
    processed_lines = []
    for line in text.split('\n'):
        if line.strip().startswith('\\'): processed_lines.append(line)
        else:
            processed_line = ""
            for char in line:
                if char in '&_#%$': processed_line += '\\' + char
                elif char == '^': processed_line += '\\textasciicircum{}'
                elif char == '~': processed_line += '\\textasciitilde{}'
                else: processed_line += char
            processed_lines.append(processed_line)
    return '\n'.join(processed_lines)

def markdown_to_html(text):
    # Process block-level elements first
    text = re.sub(r'^# (.+)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.+)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
    text = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', text, flags=re.MULTILINE)
    text = re.sub(r'^##### (.+)$', r'<h5>\1</h5>', text, flags=re.MULTILINE)
    text = re.sub(r'^###### (.+)$', r'<h6>\1</h6>', text, flags=re.MULTILINE)
    text = re.sub(r'```.*?\n(.*?)\n```', r'<pre><code>\1</code></pre>', text, flags=re.DOTALL)
    
    # Process inline elements
    text = re.sub(r'\*\*(.+?)\*\*|__(.+?)__', r'<strong>\1\2</strong>', text)
    text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)|(?<!_)_([^_]+?)_(?!_)', r'<em>\1\2</em>', text)
    text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)

    # Process paragraphs and line breaks
    html_out = []
    for paragraph in text.split('\n\n'):
        if paragraph.strip():
            # Check if paragraph is already a block element
            if any(paragraph.strip().startswith(tag) for tag in ['<h', '<pre']):
                html_out.append(paragraph)
            else:
                # Wrap in <p> and replace single newlines with <br>
                formatted_paragraph = paragraph.replace('\n', '<br>\n')
                html_out.append(f'<p>{formatted_paragraph}</p>')
    return '\n'.join(html_out)

def add_markdown_to_docx(doc, text):
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('### '): doc.add_heading(line[4:], level=3)
        else:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|(?<!\*)\*(?!\*)[^*]*?\*(?!\*)|(?<!_)_(?!_)[^_]*?_(?!_)|`[^`]*?`)', line)
            for part in parts:
                if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')): para.add_run(part[2:-2]).bold = True
                elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')): para.add_run(part[1:-1]).italic = True
                elif part.startswith('`') and part.endswith('`'): run = para.add_run(part[1:-1]); run.font.name = 'Courier New'
                else: para.add_run(part)

def split_content_by_h1(text):
    sections, lines, current_section, current_h1 = [], text.split('\n'), [], None
    for line in lines:
        if line.startswith('# '):
            if current_h1 is not None: sections.append({'title': current_h1, 'content': '\n'.join(current_section)})
            current_h1 = line[2:].strip(); current_section = [line]
        else: current_section.append(line)
    if current_h1 is not None: sections.append({'title': current_h1, 'content': '\n'.join(current_section)})
    return sections

# === Save to all formats ===

# 1. Styled HTML
html_filename = f"{output_prefix}.html"
html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Journal Entries</title>
    {HTML_CSS_STYLE}
</head>
<body>
    <div class="container">
        <div class="main-title"><h1>Journal Entries</h1></div>
"""
for note in notes:
    html_content += f"""
        <div class="entry">
            <div class="entry-date">Date: {note['date']}</div>
            <div class="entry-content">
                {markdown_to_html(note['text'])}
            </div>
        </div>"""
html_content += "</div></body></html>"
with open(html_filename, "w", encoding="utf-8") as f: f.write(html_content)

# 2. Markdown
md_filename = f"{output_prefix}.md"
md_parts = [f"## Date: {note['date']}\n\n{note['text']}" for note in notes]
with open(md_filename, "w", encoding="utf-8") as f: f.write("\n\n---\n\n".join(md_parts))

# 3. Plain Text
txt_filename = f"{output_prefix}.txt"
txt_parts = [f"Date: {note['date']}\n{markdown_to_plain_text(note['text'])}" for note in notes]
with open(txt_filename, "w", encoding="utf-8") as f: f.write("\n\n".join(txt_parts))

# 4. LaTeX
tex_filename = f"{output_prefix}.tex"
with open(tex_filename, "w", encoding="utf-8") as f:
    f.write(r"\documentclass[a4paper,12pt]{article}\n\usepackage[utf8]{inputenc}\n\usepackage{hyperref}\n\usepackage{url}\n\usepackage{lipsum}\n\usepackage{titlesec}\n\usepackage{tocloft}\n\usepackage{fancyhdr}\n\pagestyle{fancy}\n\fancyhf{}\n\rhead{\thepage}\n\begin{document}\n\n")
    f.write(r"\begin{titlepage}\n\centering\n\vspace*{5cm}\n{\Huge\bfseries The Journal \par}\n\vspace{1cm}\n{\Large by MJB \par}\n\vfill\n\end{titlepage}\n\n")
    f.write(r"\tableofcontents\n\newpage\n\n")
    for note in notes: f.write(f"\\textbf{{Date: {note['date']}}} {markdown_to_latex(note['text'])}\n\n\\vspace{{1em}}\n\n")
    f.write("\\end{document}")

# 5. DOCX
docx_filename = f"{output_prefix}.docx"
doc = Document(); doc.add_heading("Collected Notes", level=1)
for note in notes:
    doc.add_paragraph(f"Date: {note['date']}"); add_markdown_to_docx(doc, note['text']); doc.add_paragraph()
doc.save(docx_filename)

# 6. PDF (from DOCX)
try:
    convert(docx_filename, f"{output_prefix}.pdf")
    print(f"✅ PDF saved as: {os.path.basename(output_prefix)}.pdf")
except Exception as e:
    print(f"[!] PDF conversion failed: {e}")

# 7. EPUB
epub_filename = f"{output_prefix}.epub"
book = epub.EpubBook(); book.set_identifier('id123456'); book.set_title('Collected Notes'); book.set_language('en')
if os.path.exists("cover.jpg"):
    with open("cover.jpg", 'rb') as cover_file: book.set_cover("cover.jpg", cover_file.read())
    print("✅ Cover image added to EPUB")
else: print("⚠️ Cover image (cover.jpg) not found - EPUB will be created without cover")
chapters, toc_entries, chapter_counter = [], [], 1
h1_sections = [{'date': note['date'], 'title': s['title'], 'content': s['content']} for note in notes for s in split_content_by_h1(note['text'])]
if h1_sections:
    for section in h1_sections:
        chapter_filename = f'chap_{chapter_counter:02d}.xhtml'
        chapter = epub.EpubHtml(title=section['title'], file_name=chapter_filename, lang='en')
        # --- THIS IS THE CORRECTED LINE ---
        content_without_h1 = section['content'].split('\n', 1)[-1]
        chapter.content = f"<h2>{section['title']}</h2><p><strong>Date: {section['date']}</strong></p>\n{markdown_to_html(content_without_h1)}"
        book.add_item(chapter); chapters.append(chapter); toc_entries.append(epub.Link(chapter_filename, section['title'], f'chap{chapter_counter}')); chapter_counter += 1
else:
    for i, note in enumerate(notes):
        chapter_filename, chapter_title = f'chap_{i+1:02d}.xhtml', f"Entry {note['date']}"
        chapter = epub.EpubHtml(title=chapter_title, file_name=chapter_filename, lang='en')
        chapter.content = f"<h1>{chapter_title}</h1>\n{markdown_to_html(note['text'])}"
        book.add_item(chapter); chapters.append(chapter); toc_entries.append(epub.Link(chapter_filename, chapter_title, f'chap{i+1}'))
book.toc = tuple(toc_entries); book.add_item(epub.EpubNcx()); book.add_item(epub.EpubNav()); book.spine = ['nav'] + chapters
epub.write_epub(epub_filename, book)

# === Final Summary ===
print("\n✅ All files generated in folder:", folder_name)
print(f"- {os.path.basename(html_filename)} (Styled HTML)")
print(f"- {os.path.basename(md_filename)} (Markdown)")
print(f"- {os.path.basename(txt_filename)} (Plain Text)")
print(f"- {os.path.basename(tex_filename)} (LaTeX)")
print(f"- {os.path.basename(docx_filename)} (Word)")
print(f"- {os.path.basename(output_prefix)}.pdf (PDF)")
print(f"- {os.path.basename(epub_filename)} (EPUB)")

if h1_sections:
    print(f"\n📖 EPUB contains {len(h1_sections)} chapters based on H1 headings:")
    for section in h1_sections: print(f"  - {section['title']} (from {section['date']})")
else:
    print(f"\n📖 EPUB contains {len(notes)} chapters based on dates (no H1 headings found):")
    for note in notes: print(f"  - Entry {note['date']}")