import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert  # Windows-only
from ebooklib import epub
import os

# === Configure Input File ===
input_filename = "Journal.json"
title_of_output = "The Journal"

# === NEW: Configure LaTeX Lettrine (Large First Letter) ===
# Set to True to start each chapter with a large decorative letter.
# Set to False for normal paragraphs.
# NOTE: This feature will be automatically disabled for Persian documents.
USE_LETTRINE_IN_LATEX = False

# === Configure the Persian font for LaTeX output ===
PERSIAN_LATEX_FONT = "XB Niloofar"

# === CSS for the beautiful HTML output ===
HTML_CSS_STYLE = """
    body {
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
        line-height: 1.6;
        color: #333;
        background-color: #f8f9fa;
        margin: 0;
        padding: 2rem;
    }
    .container {
        max-width: 800px;
        margin: 0 auto;
        background-color: #ffffff;
        border-radius: 10px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
        overflow: hidden;
    }
    .entry {
        padding: 2rem 2.5rem;
        border-bottom: 1px solid #e9ecef;
    }
    .entry:last-child {
        border-bottom: none;
    }
    .entry-date {
        font-size: 1.1rem;
        font-weight: 600;
        color: #007bff;
        margin-bottom: 1rem;
    }
    .entry-content h1, .entry-content h2, .entry-content h3, .entry-content h4, .entry-content h5, .entry-content h6 {
        color: #495057;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
        line-height: 1.3;
    }
    .entry-content h1 { font-size: 1.8em; }
    .entry-content h2 { font-size: 1.5em; }
    .entry-content h3 { font-size: 1.25em; }
    .entry-content p {
        margin-top: 0;
        margin-bottom: 1rem;
    }
    .entry-content a {
        color: #0056b3;
        text-decoration: none;
    }
    .entry-content a:hover {
        text-decoration: underline;
    }
    .entry-content code {
        background-color: #e9ecef;
        padding: 0.2em 0.4em;
        margin: 0;
        font-size: 85%;
        border-radius: 3px;
    }
    .entry-content pre {
        background-color: #e9ecef;
        padding: 1rem;
        border-radius: 5px;
        overflow-x: auto;
    }
    .entry-content pre code {
        padding: 0;
        font-size: inherit;
        color: inherit;
        background-color: transparent;
    }
    .main-title {
        text-align: center;
        padding: 2rem;
        background-color: #007bff;
        color: white;
    }
    .main-title h1 {
        margin: 0;
        font-size: 2.5rem;
    }
"""

# === Setup output folder and prefix ===
folder_name = os.path.splitext(input_filename)[0]
os.makedirs(folder_name, exist_ok=True)
today_str = datetime.now().strftime("%Y-%m-%d")
base_filename = f"output_{os.path.basename(input_filename)}_{today_str}"
output_prefix = os.path.join(folder_name, base_filename)

# === Process notes based on file type ===
notes = []
if input_filename.lower().endswith('.json'):
    print(f"Processing Day One JSON file: {input_filename}")
    with open(input_filename, "r", encoding="utf-8") as f: data = json.load(f)
    entries = data.get("entries", [])
    if not entries:
        print("[!] Error: No 'entries' found in the JSON file."); exit()
    entries.sort(key=lambda x: x.get('creationDate', ''))
    for entry in entries:
        date_obj = datetime.strptime(entry.get("creationDate", ""), "%Y-%m-%dT%H:%M:%SZ")
        notes.append({'date': date_obj.strftime("%Y-%m-%d"), 'text': entry.get("text", "").strip()})
elif input_filename.lower().endswith('.md'):
    print(f"Processing Markdown file: {input_filename}")
    with open(input_filename, "r", encoding="utf-8") as f: md_content = f.read()
    mod_time = os.path.getmtime(input_filename)
    mod_date_obj = datetime.fromtimestamp(mod_time)
    notes.append({'date': mod_date_obj.strftime("%Y-%m-%d"), 'text': md_content.strip()})
else:
    print(f"[!] Error: Unsupported file type for '{input_filename}'. Please use a .json or .md file."); exit()

# === Persian Language Detection ===
def is_persian(text):
    return any('\u0600' <= char <= '\u06FF' for char in text)

# === MODIFIED: Check for Persian text once, for all formats that need it ===
contains_persian = any(is_persian(note['text']) for note in notes)

# === NEW: Function to apply LaTeX lettrine to a block of text ===
def apply_lettrine_to_content(text_block):
    """
    Finds the first word of the main text (skipping headings) and wraps it in a
    LaTeX \\lettrine command. This is applied per-chapter.
    """
    lines = text_block.split('\n')
    for i, line in enumerate(lines):
        stripped_line = line.strip()
        # Find the first line that isn't empty and doesn't look like a LaTeX command.
        if stripped_line and not stripped_line.startswith('\\'):
            # Found the first content line. Now, find the first word.
            words = stripped_line.split(None, 1)
            if not words: continue

            first_word_with_punct = words[0]
            # Use regex to separate the word from any trailing punctuation
            match = re.match(r'([a-zA-Z0-9]+)(\W*)', first_word_with_punct)

            # ### START OF BUG FIX ###
            # The original check prevented single-letter words like "I" or "A".
            # The new check just ensures a word was actually found.
            if not match:
                continue
            # ### END OF BUG FIX ###

            first_word_clean = match.group(1)
            trailing_punct = match.group(2)

            first_letter = first_word_clean[0]
            rest_of_word = first_word_clean[1:]

            # The user-requested lettrine command
            lettrine_cmd = f"\\lettrine[lines=2, lhang=0.33, loversize=0.3]{{{first_letter}}}{{{rest_of_word}}}"

            # Reconstruct the original line with the new command
            # This preserves leading whitespace and the rest of the line
            start_index = line.find(first_word_with_punct)
            end_index = start_index + len(first_word_with_punct)
            
            # Rebuild the rest of the line, handling the case where there is no more text
            rest_of_the_line = words[1] if len(words) > 1 else ''
            new_line = line[:start_index] + lettrine_cmd + trailing_punct + " " + rest_of_the_line
            
            lines[i] = new_line.strip()

            # We are done, we only apply lettrine once per text block.
            return '\n'.join(lines)

    # If no suitable line was found, return the text unmodified.
    return text_block

# === Markdown Processing Functions ===
def markdown_to_plain_text(text):
    text = re.sub(r'^#{1,6}\s+(.+)$', r'\1', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*|__(.+?)__', r'\1\2', text)
    text = re.sub(r'\*(.+?)\*|_(.+?)_', r'\1\2', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'```.*?\n(.*?)\n```', r'\1', text, flags=re.DOTALL)
    return text

def markdown_to_latex(text, use_persian_mode):
    # --- Existing Conversions ---
    text = re.sub(r'^# (.+)$', r'\\section{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'\\subsection{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.+)$', r'\\subsubsection{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*|__(.+?)__', r'\\textbf{\1\2}', text)
    text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)|(?<!_)_([^_]+?)_(?!_)', r'\\textit{\1\2}', text)
    text = re.sub(r'`(.+?)`', r'\\texttt{\1}', text)
    text = re.sub(r'```.*?\n(.*?)\n```', r'\\begin{verbatim}\n\1\n\\end{verbatim}', text, flags=re.DOTALL)
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'\\href{\2}{\1}', text)

    # --- NEW: Process lists (bullet points and numbered) ---
    lines = text.split('\n')
    processed_lines = []
    in_itemize = False
    in_enumerate = False

    for line in lines:
        # Check for numbered list items (e.g., "1. item")
        numbered_match = re.match(r'^\s*(\d+)\.\s+(.*)', line)
        # Check for bullet list items (e.g., "* item", "- item")
        bullet_match = re.match(r'^\s*([*+-])\s+(.*)', line)

        # Handle numbered lists
        if numbered_match:
            if not in_enumerate:
                processed_lines.append(r'\begin{enumerate}')
                in_enumerate = True
            processed_lines.append(r'  \item ' + numbered_match.group(2))
        else:
            if in_enumerate:
                processed_lines.append(r'\end{enumerate}')
                in_enumerate = False
            
            # Handle bullet lists
            if bullet_match:
                if not in_itemize:
                    processed_lines.append(r'\begin{itemize}')
                    in_itemize = True
                processed_lines.append(r'  \item ' + bullet_match.group(2))
            else:
                if in_itemize:
                    processed_lines.append(r'\end{itemize}')
                    in_itemize = False
                
                # If not in any list, add the line
                if not in_enumerate and not in_itemize:
                    processed_lines.append(line)

    # Close any open list environments at the end of the text
    if in_itemize:
        processed_lines.append(r'\end{itemize}')
    if in_enumerate:
        processed_lines.append(r'\end{enumerate}')
    
    text = '\n'.join(processed_lines)


    # --- Existing Character Escaping for non-Persian mode ---
    if use_persian_mode:
        return text
    else:
        # ### START OF MODIFIED BLOCK ###
        processed_lines = []
        for line in text.split('\n'):
            # The original check was insufficient. We will apply a more careful escaping
            # to all lines, but we will not escape characters that form LaTeX commands.
            if line.strip().startswith('\\begin{verbatim}'):
                 processed_lines.append(line)
                 continue

            processed_line = ""
            # This dictionary is modified to NOT escape \, {, and } which are
            # part of the LaTeX commands we generate (e.g., \textbf{}).
            special_chars = {'&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_',
                             '~': r'\textasciitilde{}', '^': r'\textasciicircum{}'}
            
            # The old check `if line.strip().startswith...` is removed and this logic applies to all lines.
            # This fixes commands that appear in the middle of a line.
            # We iterate character by character to build the escaped line.
            for char in line:
                processed_line += special_chars.get(char, char)
            processed_lines.append(processed_line)
        return '\n'.join(processed_lines)
        # ### END OF MODIFIED BLOCK ###

def markdown_to_html(text):
    text = re.sub(r'^# (.+)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.+)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
    text = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', text, flags=re.MULTILINE)
    text = re.sub(r'^##### (.+)$', r'<h5>\1</h5>', text, flags=re.MULTILINE)
    text = re.sub(r'^###### (.+)$', r'<h6>\1</h6>', text, flags=re.MULTILINE)
    text = re.sub(r'```.*?\n(.*?)\n```', r'<pre><code>\1</code></pre>', text, flags=re.DOTALL)
    text = re.sub(r'\*\*(.+?)\*\*|__(.+?)__', r'<strong>\1\2</strong>', text)
    text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)|(?<!_)_([^_]+?)_(?!_)', r'<em>\1\2</em>', text)
    text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)
    html_out = []
    for paragraph in text.split('\n\n'):
        if not paragraph.strip(): continue
        if any(paragraph.strip().startswith(tag) for tag in ['<h', '<pre']):
            html_out.append(paragraph)
        else:
            html_out.append(f'<p>{paragraph.replace("\n", "<br>")}</p>')
    return '\n'.join(html_out)

def add_markdown_to_docx(doc, text):
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('### '): doc.add_heading(line[4:], level=3)
        else:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|(?<!\*)\*(?!\*)[^*]*?\*(?!\*)|(?<!_)_(?!_)[^_]*?_(?!_)|`[^`]*?`)', line)
            for part in parts:
                if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')): para.add_run(part[2:-2]).bold = True
                elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')): para.add_run(part[1:-1]).italic = True
                elif part.startswith('`') and part.endswith('`'): run = para.add_run(part[1:-1]); run.font.name = 'Courier New'
                else: para.add_run(part)

def split_content_by_h1(text):
    sections, lines, current_section, current_h1 = [], text.split('\n'), [], None
    for line in lines:
        if line.startswith('# '):
            if current_h1 is not None: sections.append({'title': current_h1, 'content': '\n'.join(current_section)})
            current_h1 = line[2:].strip(); current_section = [line]
        else: current_section.append(line)
    if current_h1 is not None: sections.append({'title': current_h1, 'content': '\n'.join(current_section)})
    if not sections:
        return [{'title': None, 'content': text}]
    return sections

# === Save to all formats ===

h1_sections = [{'date': note['date'], 'title': s['title'], 'content': s['content']} for note in notes for s in split_content_by_h1(note['text'])]

# === 1. Styled HTML (MODIFIED to add RTL conditionally) ===
html_filename = f"{output_prefix}.html"
final_css = HTML_CSS_STYLE
html_attrs = 'lang="en"'

if contains_persian:
    print("Persian text detected. Applying RTL direction to HTML output.")
    html_attrs = 'lang="fa" dir="rtl"'
    # Add CSS rules needed for proper RTL display
    rtl_css = """
    /* --- RTL SUPPORT --- */
    html[dir="rtl"] body {
        direction: rtl;
        text-align: right;
    }
    /* Keep code blocks LTR for correct formatting */
    html[dir="rtl"] pre,
    html[dir="rtl"] code {
        direction: ltr;
        text-align: left;
    }
    """
    final_css += rtl_css

with open(html_filename, "w", encoding="utf-8") as f:
    f.write(f'<!DOCTYPE html><html {html_attrs}><head><meta charset="UTF-8"><title>{title_of_output}</title><style>{final_css}</style></head><body>\n')
    f.write(f'<div class="container"><div class="main-title"><h1>{title_of_output}</h1></div>\n')
    for note in notes:
        f.write(f'<div class="entry"><div class="entry-date">Date: {note["date"]}</div>\n')
        f.write(f'<div class="entry-content">{markdown_to_html(note["text"])}</div></div>\n')
    f.write('</div></body></html>')


# 2. Markdown
md_filename = f"{output_prefix}.md"
with open(md_filename, "w", encoding="utf-8") as f:
    md_parts = [f"## Date: {note['date']}\n\n{note['text']}" for note in notes]
    f.write("\n\n---\n\n".join(md_parts))

# 3. Plain Text
txt_filename = f"{output_prefix}.txt"
with open(txt_filename, "w", encoding="utf-8") as f:
    txt_parts = [f"Date: {note['date']}\n{markdown_to_plain_text(note['text'])}" for note in notes]
    f.write("\n\n".join(txt_parts))

# 4. LaTeX
tex_filename = f"{output_prefix}.tex"
with open(tex_filename, "w", encoding="utf-8") as f:
    preamble = [
        r"\documentclass[a4paper,12pt]{article}",
        r"\usepackage{fontspec}",
        r"\setmainfont{Alice}[AutoFakeBold=2.0]",
        r"\usepackage{hyperref}",
        r"\usepackage{fancyhdr}",
        r"\usepackage{graphicx}",
        r"\usepackage{setspace}",
        r"\setlength{\headheight}{15pt}"
    ]
    # MODIFICATION: Add the lettrine package ONLY if enabled AND text is not Persian.
    lettrine_is_active = USE_LETTRINE_IN_LATEX and not contains_persian
    if lettrine_is_active:
        print("Lettrine feature is enabled for LaTeX (English text only).")
        preamble.append(r"\usepackage{lettrine}")

    if contains_persian:
        print(f"Persian text detected. Using XePersian with font '{PERSIAN_LATEX_FONT}' for LaTeX output.")
        if USE_LETTRINE_IN_LATEX:
            print(" -> Lettrine feature disabled for Persian text.")
        preamble.append(r"\usepackage{xepersian}")
        preamble.append(f"\\settextfont{{{PERSIAN_LATEX_FONT}}}")
    else:
        print("No Persian text detected. Using standard LaTeX output.")
        preamble.append(r"\usepackage[utf8]{inputenc}")
    
    f.write("\n".join(preamble) + "\n")
    f.write(r"\hypersetup{colorlinks=true, linkcolor=blue, urlcolor=blue, pdfproducer={Python Script}, pdftitle={Collected Notes}}" + "\n")
    f.write("\\pagestyle{fancy}\n\\fancyhf{}\n\\rhead{\\thepage}\n")
    f.write("\\begin{document}" + "\n\n")
    f.write("\\onehalfspacing" + "\n\n")
    f.write("\\begin{titlepage}\n\\centering\n\\vspace*{5cm}\n{\\Huge\\bfseries")
    f.write(f" {title_of_output} ")
    f.write("\\par}\n\\vfill\n\\end{titlepage}" + "\n\n")
    f.write(r"\tableofcontents" + "\n" + r"\newpage" + "\n\n")
    
    has_titles = h1_sections and h1_sections[0]['title'] is not None

    if has_titles:
        print("Found H1 headings. Using titles for LaTeX chapters.")
        for section in h1_sections:
            processed_text = markdown_to_latex(section['content'], contains_persian)
            # MODIFICATION: Apply lettrine only if the flag is active.
            if lettrine_is_active:
                processed_text = apply_lettrine_to_content(processed_text)
            f.write(f"{processed_text}\n\n\\newpage\n\n")
    else:
        print("No H1 headings found. Using dates for LaTeX chapters.")
        for note in notes:
            f.write(f"\\section{{Entry: {note['date']}}}\n")
            processed_text = markdown_to_latex(note['text'], contains_persian)
            # MODIFICATION: Apply lettrine only if the flag is active.
            if lettrine_is_active:
                processed_text = apply_lettrine_to_content(processed_text)
            f.write(f"{processed_text}\n\n\\newpage\n\n")
        
    f.write("\\end{document}")

# 5. DOCX
docx_filename = f"{output_prefix}.docx"
doc = Document(); doc.add_heading(title_of_output, level=1)
for note in notes:
    doc.add_heading(f"Date: {note['date']}", level=2); add_markdown_to_docx(doc, note['text']); doc.add_paragraph()
doc.save(docx_filename)

# 6. PDF (from DOCX)
try:
    print("Attempting to convert DOCX to PDF...")
    convert(docx_filename, f"{output_prefix}.pdf")
    print(f"✅ PDF saved as: {os.path.basename(output_prefix)}.pdf")
except Exception as e:
    print(f"[!] PDF conversion from DOCX failed. This is a Windows-only feature and requires MS Word.")
    print(f"    To generate a PDF from the LaTeX file, run 'xelatex \"{os.path.basename(tex_filename)}\"' in your terminal.")
    print(f"    Error details: {e}")

# 7. EPUB
epub_filename = f"{output_prefix}.epub"
book = epub.EpubBook(); book.set_identifier('id123456'); book.set_title(title_of_output); book.set_language('en')
if os.path.exists("cover.jpg"):
    book.set_cover("cover.jpg", open("cover.jpg", 'rb').read())
    print("✅ Cover image added to EPUB")
else: print("⚠️ Cover image (cover.jpg) not found - EPUB will be created without cover")
chapters, toc_entries = [], []

has_titles_epub = h1_sections and h1_sections[0]['title'] is not None
if has_titles_epub:
    for i, section in enumerate(h1_sections):
        chapter_filename = f'chap_{i+1:02d}.xhtml'
        chapter_title = section['title']
        content_without_h1 = section['content'].split('\n', 1)[-1]
        chapter = epub.EpubHtml(title=chapter_title, file_name=chapter_filename, lang='en')
        chapter.content = f"<h1>{chapter_title}</h1><p><strong>Date: {section['date']}</strong></p>\n{markdown_to_html(content_without_h1)}"
        book.add_item(chapter); chapters.append(chapter); toc_entries.append(epub.Link(chapter_filename, chapter_title, f'chap{i+1}'))
else:
    for i, note in enumerate(notes):
        chapter_filename, chapter_title = f'chap_{i+1:02d}.xhtml', f"Entry {note['date']}"
        chapter = epub.EpubHtml(title=chapter_title, file_name=chapter_filename, lang='en')
        chapter.content = f"<h1>{chapter_title}</h1>\n{markdown_to_html(note['text'])}"
        book.add_item(chapter); chapters.append(chapter); toc_entries.append(epub.Link(chapter_filename, chapter_title, f'chap{i+1}'))
book.toc = tuple(toc_entries); book.add_item(epub.EpubNcx()); book.add_item(epub.EpubNav()); book.spine = ['nav'] + chapters
epub.write_epub(epub_filename, book)

# === Final Summary ===
print("\n✅ All files generated in folder:", folder_name)
print(f"- {os.path.basename(html_filename)} (Styled HTML)")
print(f"- {os.path.basename(md_filename)} (Markdown)")
print(f"- {os.path.basename(txt_filename)} (Plain Text)")
print(f"- {os.path.basename(tex_filename)} (LaTeX)")
print(f"- {os.path.basename(docx_filename)} (Word)")
if os.path.exists(f"{output_prefix}.pdf"): print(f"- {os.path.basename(output_prefix)}.pdf (PDF)")
print(f"- {os.path.basename(epub_filename)} (EPUB)")

if has_titles_epub:
    print(f"\n📖 EPUB and LaTeX contain {len(h1_sections)} chapters based on H1 headings:")
    for section in h1_sections: print(f"  - {section['title']} (from {section['date']})")
else:
    print(f"\n📖 EPUB and LaTeX contain {len(notes)} chapters based on dates (no H1 headings found):")
    for note in notes: print(f"  - Entry {note['date']}")