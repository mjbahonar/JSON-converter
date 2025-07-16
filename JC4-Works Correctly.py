import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert  # Windows-only
from ebooklib import epub
import os

# === Load JSON data ===
jsonName="Journal-25-june.json"
with open(jsonName, "r", encoding="utf-8") as f:
    data = json.load(f)

# === Setup output prefix ===
today_str = datetime.now().strftime("%Y-%m-%d")
output_prefix = f"output_{jsonName}_{today_str}"

# === Process notes with creation date ===
entries = data["entries"]
notes = []
for entry in entries:
    creation_date_str = entry.get("creationDate", "")
    date_obj = datetime.strptime(creation_date_str, "%Y-%m-%dT%H:%M:%SZ")
    formatted_date = date_obj.strftime("%Y-%m-%d")
    note_text = entry["text"].strip()
    notes.append({
        'date': formatted_date,
        'text': note_text
    })

# === Markdown Processing Functions ===
def markdown_to_plain_text(text):
    """Convert markdown to plain text"""
    # Headers
    text = re.sub(r'^#{1,6}\s+(.+)$', r'\1', text, flags=re.MULTILINE)
    # Bold and italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Links
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # Code
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'```.*?\n(.*?)\n```', r'\1', text, flags=re.DOTALL)
    return text

def markdown_to_latex(text):
    """Convert markdown to LaTeX with proper formatting"""
    
    # First, handle markdown formatting before escaping
    # Headers
    text = re.sub(r'^# (.+)$', r'\\section{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'\\subsection{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.+)$', r'\\subsubsection{\1}', text, flags=re.MULTILINE)
    text = re.sub(r'^#{4,6} (.+)$', r'\\paragraph{\1}', text, flags=re.MULTILINE)
    
    # Bold formatting - handle both ** and __
    text = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', text)
    text = re.sub(r'__(.+?)__', r'\\textbf{\1}', text)
    
    # Italic formatting - handle both * and _
    text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'\\textit{\1}', text)
    text = re.sub(r'(?<!_)_([^_]+?)_(?!_)', r'\\textit{\1}', text)
    
    # Code formatting
    text = re.sub(r'`(.+?)`', r'\\texttt{\1}', text)
    text = re.sub(r'```.*?\n(.*?)\n```', r'\\begin{verbatim}\n\1\n\\end{verbatim}', text, flags=re.DOTALL)
    
    # Links
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'\\href{\2}{\1}', text)
    
    # Now escape LaTeX special characters, but preserve our LaTeX commands
    # Split text into lines and process each line
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        # Skip lines that are LaTeX commands
        if line.strip().startswith('\\'):
            processed_lines.append(line)
        else:
            # Process the line for escaping, but be careful with our LaTeX commands
            processed_line = ""
            i = 0
            while i < len(line):
                if line[i:i+1] == '\\' and i < len(line) - 1:
                    # Find the end of this LaTeX command
                    j = i + 1
                    while j < len(line) and (line[j].isalpha() or line[j] in '{}'):
                        j += 1
                        if j < len(line) and line[j-1] == '}':
                            break
                    # Add the LaTeX command as-is
                    processed_line += line[i:j]
                    i = j
                else:
                    # Escape individual characters that need escaping
                    char = line[i]
                    if char in '&_#%$':
                        processed_line += '\\' + char
                    elif char == '^':
                        processed_line += '\\textasciicircum{}'
                    elif char == '~':
                        processed_line += '\\textasciitilde{}'
                    else:
                        processed_line += char
                    i += 1
            processed_lines.append(processed_line)
    
    return '\n'.join(processed_lines)

def markdown_to_html(text):
    """Convert markdown to HTML"""
    # Headers
    text = re.sub(r'^# (.+)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.+)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
    text = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', text, flags=re.MULTILINE)
    text = re.sub(r'^##### (.+)$', r'<h5>\1</h5>', text, flags=re.MULTILINE)
    text = re.sub(r'^###### (.+)$', r'<h6>\1</h6>', text, flags=re.MULTILINE)
    
    # Bold and italic
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'__(.+?)__', r'<strong>\1</strong>', text)
    text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
    text = re.sub(r'_(.+?)_', r'<em>\1</em>', text)
    
    # Code
    text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
    text = re.sub(r'```.*?\n(.*?)\n```', r'<pre><code>\1</code></pre>', text, flags=re.DOTALL)
    
    # Links
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)
    
    # Line breaks
    text = re.sub(r'\n\n', '</p><p>', text)
    text = re.sub(r'\n', '<br>', text)
    
    return f'<p>{text}</p>'

def add_markdown_to_docx(doc, text):
    """Add markdown text to Word document with proper formatting"""
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
            
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            heading = doc.add_heading(line[6:], level=5)
        elif line.startswith('###### '):
            heading = doc.add_heading(line[7:], level=6)
        else:
            # Regular paragraph - handle bold and italic
            para = doc.add_paragraph()
            
            # Split by formatting markers
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|(?<!\*)\*(?!\*)[^*]*?\*(?!\*)|(?<!_)_(?!_)[^_]*?_(?!_)|`[^`]*?`)', line)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    para.add_run(part[2:-2]).bold = True
                elif part.startswith('__') and part.endswith('__'):
                    para.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    para.add_run(part[1:-1]).italic = True
                elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
                    para.add_run(part[1:-1]).italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    para.add_run(part)
        i += 1

def extract_h1_titles(text):
    """Extract all H1 titles from markdown text"""
    h1_matches = re.findall(r'^# (.+)$', text, flags=re.MULTILINE)
    return h1_matches

def split_content_by_h1(text):
    """Split content into sections based on H1 headings"""
    sections = []
    lines = text.split('\n')
    current_section = []
    current_h1 = None
    
    for line in lines:
        if line.startswith('# '):
            # If we have a previous section, save it
            if current_h1 is not None:
                sections.append({
                    'title': current_h1,
                    'content': '\n'.join(current_section)
                })
            
            # Start new section
            current_h1 = line[2:].strip()
            current_section = [line]
        else:
            current_section.append(line)
    
    # Don't forget the last section
    if current_h1 is not None:
        sections.append({
            'title': current_h1,
            'content': '\n'.join(current_section)
        })
    
    return sections

# === Save as TXT (plain text) ===
plain_text_notes = []
for note in notes:
    plain_content = markdown_to_plain_text(note['text'])
    plain_text_notes.append(f"Date: {note['date']}\n{plain_content}")

full_plain_text = "\n\n".join(plain_text_notes)
with open(f"{output_prefix}.txt", "w", encoding="utf-8") as f:
    f.write(full_plain_text)

# === Save as LaTeX ===
with open(f"{output_prefix}.tex", "w", encoding="utf-8") as f:
    f.write(r"\documentclass[a4paper,12pt]{article}" + "\n")
    f.write(r"\usepackage[utf8]{inputenc}" + "\n")
    f.write(r"\usepackage{hyperref}" + "\n")
    f.write(r"\usepackage{url}" + "\n")
    f.write(r"\usepackage{lipsum}" + "\n")
    f.write(r"\usepackage{titlesec}" + "\n")
    f.write(r"\usepackage{tocloft}" + "\n")
    f.write(r"\usepackage{fancyhdr}" + "\n")
    f.write(r"\pagestyle{fancy}" + "\n")
    f.write(r"\fancyhf{}" + "\n")
    f.write(r"\rhead{\thepage}" + "\n")
    f.write(r"\begin{document}" + "\n\n")

    # Cover Page
    f.write(r"\begin{titlepage}" + "\n")
    f.write(r"\centering" + "\n")
    f.write(r"\vspace*{5cm}" + "\n")
    f.write(r"{\Huge\bfseries The Journal \par}" + "\n")
    f.write(r"\vspace{1cm}" + "\n")
    f.write(r"{\Large by MJB \par}" + "\n")
    f.write(r"\vfill" + "\n")
    f.write(r"\end{titlepage}" + "\n\n")

    # Table of contents
    f.write(r"\tableofcontents" + "\n")
    f.write(r"\newpage" + "\n\n")

    # Notes content
    for i, note in enumerate(notes):
        f.write(f"\\textbf{{Date: {note['date']}}} ")
        #f.write(r"\addcontentsline{toc}{section}{" + f"Date: {note['date']}" + "}\n")
        latex_content = markdown_to_latex(note['text'])
        f.write(latex_content + "\n\n")
        f.write("\\vspace{1em}\n\n")

    f.write("\\end{document}")


# === Save as DOCX ===
docx_filename = f"{output_prefix}.docx"
doc = Document()
doc.add_heading("Collected Notes", level=1)

for note in notes:
    # Add date as a subheading
    doc.add_paragraph(f"Date: {note['date']}")
    
    # Add formatted content
    add_markdown_to_docx(doc, note['text'])
    
    # Add some space between notes
    doc.add_paragraph()

doc.save(docx_filename)

# === Convert DOCX to PDF (Windows with MS Word) ===
try:
    convert(docx_filename, f"{output_prefix}.pdf")
    print(f"✅ PDF saved as: {output_prefix}.pdf")
except Exception as e:
    print(f"[!] PDF conversion failed: {e}")

# === Save as EPUB with H1 headings in TOC ===
book = epub.EpubBook()
book.set_identifier('id123456')
book.set_title('Collected Notes')
book.set_language('en')

# Add cover image if it exists
cover_path = "cover.jpg"
if os.path.exists(cover_path):
    with open(cover_path, 'rb') as cover_file:
        book.set_cover("cover.jpg", cover_file.read())
    print("✅ Cover image added to EPUB")
else:
    print("⚠️ Cover image (cover.jpg) not found - EPUB will be created without cover")

chapters = []
toc_entries = []
chapter_counter = 1

# Collect all content and extract H1 sections
all_content = ""
for note in notes:
    all_content += f"Date: {note['date']}\n{note['text']}\n\n"

# Check if there are any H1 headings in the content
h1_sections = []
for note in notes:
    sections = split_content_by_h1(note['text'])
    for section in sections:
        h1_sections.append({
            'date': note['date'],
            'title': section['title'],
            'content': section['content']
        })

if h1_sections:
    # Create chapters based on H1 headings
    for section in h1_sections:
        chapter_filename = f'chap_{chapter_counter:02d}.xhtml'
        chapter = epub.EpubHtml(title=section['title'], file_name=chapter_filename, lang='en')
        
        # Create HTML content for this chapter
        html_content = f"<p><strong>Date: {section['date']}</strong></p>\n"
        
        # Convert the section content to HTML (this already includes the H1)
        section_html = markdown_to_html(section['content'])
        html_content += section_html
        
        chapter.content = html_content
        book.add_item(chapter)
        chapters.append(chapter)
        
        # Add to table of contents
        toc_entries.append(epub.Link(chapter_filename, section['title'], f'chap{chapter_counter}'))
        chapter_counter += 1
    
    # Set up table of contents
    book.toc = tuple(toc_entries)
    
else:
    # No H1 headings found, create chapters by date
    for i, note in enumerate(notes):
        chapter_filename = f'chap_{i+1:02d}.xhtml'
        chapter_title = f"Entry {note['date']}"
        chapter = epub.EpubHtml(title=chapter_title, file_name=chapter_filename, lang='en')
        
        html_content = f"<h1>{chapter_title}</h1>\n"
        html_content += markdown_to_html(note['text'])
        
        chapter.content = html_content
        book.add_item(chapter)
        chapters.append(chapter)
        
        # Add to table of contents
        toc_entries.append(epub.Link(chapter_filename, chapter_title, f'chap{i+1}'))
    
    # Set up table of contents
    book.toc = tuple(toc_entries)

# Add required EPUB components
book.add_item(epub.EpubNcx())
book.add_item(epub.EpubNav())
book.spine = ['nav'] + chapters

# Write the EPUB file
epub.write_epub(f"{output_prefix}.epub", book)

print("✅ All files generated:")
print(f"- {output_prefix}.txt (plain text)")
print(f"- {output_prefix}.tex (LaTeX with formatting)")
print(f"- {output_prefix}.docx (Word with formatting)")
print(f"- {output_prefix}.pdf (from Word)")
print(f"- {output_prefix}.epub (eBook with H1 headings in TOC)")

# Print summary of EPUB structure
if h1_sections:
    print(f"\n📖 EPUB contains {len(h1_sections)} chapters based on H1 headings:")
    for section in h1_sections:
        print(f"  - {section['title']} (from {section['date']})")
else:
    print(f"\n📖 EPUB contains {len(notes)} chapters based on dates (no H1 headings found):")